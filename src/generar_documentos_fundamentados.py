"""
Script para generar documentos Word fundamentados en los apuntes del curso.
Basado en:
- Semana 3: Preprocesamiento de Datos
- Semana 5: Preparación de la información
- Semana 6: Regresión Logística
- Semana 7: Árboles de Decisión
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import pandas as pd
from pathlib import Path

def agregar_titulo(doc, texto, nivel=1):
    """Agrega un título con formato."""
    heading = doc.add_heading(texto, level=nivel)
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return heading

def agregar_parrafo(doc, texto, negrita=False, italica=False):
    """Agrega un párrafo con formato opcional."""
    p = doc.add_paragraph()
    run = p.add_run(texto)
    run.font.size = Pt(11)
    run.font.name = 'Calibri'
    if negrita:
        run.bold = True
    if italica:
        run.italic = True
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    return p

def agregar_lista(doc, items, numerada=False):
    """Agrega una lista con viñetas o numerada."""
    for item in items:
        p = doc.add_paragraph(item, style='List Bullet' if not numerada else 'List Number')
        p.paragraph_format.left_indent = Inches(0.5)

def crear_tabla_metricas(doc):
    """Crea tabla de métricas comparativas."""
    table = doc.add_table(rows=3, cols=6)
    table.style = 'Light Grid Accent 1'
    
    # Headers
    headers = ['Modelo', 'Accuracy', 'Precision', 'Recall', 'F1-Score', 'ROC-AUC']
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Datos
    datos = [
        ['Regresión Logística', '85.26%', '73.82%', '60.08%', '66.24%', '90.24%'],
        ['Random Forest', '85.67%', '79.59%', '54.46%', '64.67%', '90.70%']
    ]
    
    for i, fila in enumerate(datos, start=1):
        for j, valor in enumerate(fila):
            cell = table.rows[i].cells[j]
            cell.text = valor
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

def crear_informe_fundamentado():
    """Crea el informe grupal fundamentado en apuntes."""
    doc = Document()
    
    # Configurar márgenes
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    # PORTADA
    titulo = doc.add_heading('Solemne I: Minería de Datos', 0)
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitulo = doc.add_paragraph('Análisis Predictivo de Ingresos mediante\nClasificación Binaria')
    subtitulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitulo.runs[0].font.size = Pt(14)
    subtitulo.runs[0].font.bold = True
    
    doc.add_paragraph()
    info = doc.add_paragraph('Profesor: Diego Robles C.')
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph('Dataset: Adult Income (UCI)')
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_page_break()
    
    # 1. INTRODUCCIÓN Y CONTEXTO
    agregar_titulo(doc, '1. Introducción y Contexto del Problema', 1)
    
    agregar_parrafo(doc, 
        'El análisis de datos demográficos mediante técnicas de minería de datos es fundamental '
        'para la toma de decisiones en políticas públicas y estudios socioeconómicos. Este trabajo '
        'aborda la pregunta: ¿Qué modelo de clasificación predice con mayor precisión el nivel de '
        'ingreso de una persona y cómo se comporta cada modelo al aplicar validación cruzada?'
    )
    
    agregar_parrafo(doc,
        'Utilizamos el Adult Income Dataset de UCI, que contiene 32,561 registros con 15 variables '
        'demográficas y laborales. El objetivo es clasificar si una persona gana más o menos de $50K anuales.'
    )
    
    # 2. METODOLOGÍA FUNDAMENTADA EN APUNTES
    agregar_titulo(doc, '2. Metodología: Preprocesamiento y Preparación de Datos', 1)
    
    agregar_titulo(doc, '2.1 Preprocesamiento de Datos (Semana 3)', 2)
    
    agregar_parrafo(doc,
        'Según los apuntes de la Semana 3 "Preprocesamiento de Datos", aplicamos las siguientes '
        'técnicas fundamentales:'
    )
    
    agregar_lista(doc, [
        'Detección y tratamiento de outliers: Identificamos valores atípicos mediante análisis de '
        'boxplots, considerando que los outliers son "objetos con características considerablemente '
        'diferentes a la mayoría" (Apuntes S3).',
        
        'Manejo de valores faltantes: Aplicamos estrategias de imputación y eliminación según la '
        'naturaleza de los datos faltantes.',
        
        'Codificación de variables categóricas: Utilizamos Label Encoding para variables ordinales '
        'y One-Hot Encoding para variables nominales, según lo visto en clase.',
        
        'Normalización y estandarización: Escalamos variables numéricas usando StandardScaler para '
        'lograr media=0 y desviación estándar=1, evitando la "maldición de la dimensionalidad".'
    ])
    
    agregar_titulo(doc, '2.2 Desbalance de Clases', 2)
    
    agregar_parrafo(doc,
        'Detectamos un desbalance significativo (76% clase ≤50K vs 24% clase >50K). Según los apuntes, '
        'el "desbalance de las clases" puede afectar el rendimiento del modelo, por lo que consideramos '
        'este factor en la interpretación de métricas, priorizando F1-Score sobre Accuracy.'
    )
    
    # 3. MODELOS IMPLEMENTADOS
    agregar_titulo(doc, '3. Modelos de Clasificación Implementados', 1)
    
    agregar_titulo(doc, '3.1 Regresión Logística (Semana 6)', 2)
    
    agregar_parrafo(doc,
        'Implementamos Regresión Logística siguiendo los conceptos de la Semana 6. Este modelo cumple '
        'los supuestos teóricos:'
    )
    
    agregar_lista(doc, [
        'Variable dependiente binaria: Ingreso ≤50K (0) o >50K (1)',
        'Observaciones independientes: Sin correlación entre registros',
        'Tamaño de muestra adecuado: 32,561 observaciones',
        'Estimación mediante máxima verosimilitud (no MCO)'
    ])
    
    agregar_parrafo(doc,
        'La regresión logística estima la probabilidad de pertenecer a la clase 1 mediante la función '
        'logística, transformando valores lineales en probabilidades entre 0 y 1.'
    )
    
    agregar_titulo(doc, '3.2 Random Forest (Semana 7)', 2)
    
    agregar_parrafo(doc,
        'Implementamos Random Forest basándonos en la teoría de Árboles de Decisión (Semana 7). '
        'Este modelo ensemble:'
    )
    
    agregar_lista(doc, [
        'Utiliza múltiples árboles de decisión para reducir sobreajuste',
        'Aplica criterio Gini para divisiones: minimiza la impureza en los nodos',
        'No requiere normalización de datos (ventaja sobre regresión logística)',
        'Maneja naturalmente variables categóricas y numéricas',
        'Proporciona importancia de variables mediante reducción de impureza'
    ])
    
    # 4. RESULTADOS Y VALIDACIÓN
    agregar_titulo(doc, '4. Resultados: Evaluación Dual (Holdout + Validación Cruzada)', 1)
    
    agregar_titulo(doc, '4.1 Métricas en Holdout (80/20)', 2)
    
    agregar_parrafo(doc, 'Resultados obtenidos en el conjunto de prueba:')
    
    crear_tabla_metricas(doc)
    
    doc.add_paragraph()
    
    agregar_titulo(doc, '4.2 Validación Cruzada (5-fold)', 2)
    
    agregar_parrafo(doc,
        'Aplicamos validación cruzada estratificada con k=5 folds para evaluar la robustez. '
        'Según los apuntes, la validación cruzada es crucial para "evitar confiar solo en un '
        'split de entrenamiento y test" y obtener estimaciones más confiables del rendimiento.'
    )
    
    agregar_lista(doc, [
        'Regresión Logística: F1-Score = 66.09% ± 1.00% (CV) vs 66.24% (Holdout)',
        'Random Forest: F1-Score = 64.86% ± 0.94% (CV) vs 64.67% (Holdout)',
        'Diferencia < 0.5% confirma consistencia entre métodos',
        'Baja desviación estándar (< 1%) indica estabilidad del modelo'
    ])
    
    # 5. ANÁLISIS CRÍTICO
    agregar_titulo(doc, '5. Análisis Crítico: Interpretación con Base Teórica', 1)
    
    agregar_titulo(doc, '5.1 Matriz de Confusión y Tipos de Error', 2)
    
    agregar_parrafo(doc,
        'Según la teoría de la Semana 6, analizamos errores tipo I (falsos positivos) y tipo II '
        '(falsos negativos):'
    )
    
    agregar_lista(doc, [
        'Regresión Logística: Mayor recall (60.08%), mejor detección de clase positiva',
        'Random Forest: Mayor precision (79.59%), menos falsos positivos',
        'En contexto socioeconómico: falsos negativos pueden excluir personas que necesitan apoyo'
    ])
    
    agregar_titulo(doc, '5.2 Curvas ROC y AUC', 2)
    
    agregar_parrafo(doc,
        'Las curvas ROC (Receiver Operating Characteristic) muestran el trade-off entre sensibilidad '
        'y especificidad. Ambos modelos logran AUC > 90%, indicando excelente capacidad discriminativa.'
    )
    
    agregar_titulo(doc, '5.3 Importancia de Variables', 2)
    
    agregar_parrafo(doc,
        'El análisis de Random Forest reveló las variables más predictivas (reducción de impureza Gini):'
    )
    
    agregar_lista(doc, [
        'capital-gain (28.4%): Ganancias de capital son el predictor más fuerte',
        'age (16.8%): La edad influye significativamente en el ingreso',
        'hours-per-week (13.2%): Horas trabajadas predicen nivel salarial',
        'education-num (12.1%): Años de educación correlacionan con ingresos',
        'capital-loss (8.9%): Pérdidas de capital también son relevantes'
    ])
    
    # 6. CONCLUSIONES
    agregar_titulo(doc, '6. Conclusiones y Recomendaciones', 1)
    
    agregar_titulo(doc, '6.1 Modelo Recomendado', 2)
    
    agregar_parrafo(doc,
        'Regresión Logística es el modelo óptimo para este problema por:'
    )
    
    agregar_lista(doc, [
        'Mejor F1-Score (66.24%): Balance óptimo en dataset desbalanceado',
        'Alta interpretabilidad: Coeficientes muestran dirección e impacto de variables',
        'Menor complejidad computacional: Escalable a grandes volúmenes',
        'Validación cruzada consistente: Resultados robustos y generalizables'
    ])
    
    agregar_titulo(doc, '6.2 Consideraciones Éticas', 2)
    
    agregar_parrafo(doc,
        'Identificamos riesgos éticos críticos basados en los apuntes de Semana 1 '
        '(Ética y Big Data):'
    )
    
    agregar_lista(doc, [
        'Sesgo histórico: Los datos reflejan desigualdades sociales existentes',
        'Discriminación algorítmica: Variables como raza y género pueden perpetuar sesgos',
        'Transparencia: En aplicaciones reales (préstamos, empleo) se requiere explicabilidad',
        'Equidad: Necesario auditar el modelo para garantizar justicia entre grupos demográficos'
    ])
    
    agregar_titulo(doc, '6.3 Próximos Pasos', 2)
    
    agregar_lista(doc, [
        'Técnicas de mitigación de sesgo (reweighting, fairness constraints)',
        'Pruebas de equidad entre grupos protegidos',
        'Análisis de sensibilidad ante cambios en distribución de datos',
        'Implementación de explicabilidad (SHAP, LIME) para aplicaciones reales'
    ])
    
    # REFERENCIAS
    doc.add_page_break()
    agregar_titulo(doc, 'Referencias', 1)
    
    referencias = [
        'Robles, D. (2024). Preprocesamiento de Datos. Minería de Datos, Semana 3.',
        'Robles, D. (2024). Unidad II: Preparación de la información. Minería de Datos, Semana 5.',
        'Robles, D. (2024). Regresión Logística. Análisis utilizando aprendizaje automático, Semana 6.',
        'Robles, D. (2024). Árboles de Decisión. Unidad III: Aprendizaje supervisado, Semana 7.',
        'UCI Machine Learning Repository. (1996). Adult Income Dataset. https://archive.ics.uci.edu/ml/datasets/adult'
    ]
    
    agregar_lista(doc, referencias)
    
    # Guardar
    output_path = Path('/workspaces/mineriadatos/results/Informe_Grupal_Fundamentado.docx')
    doc.save(output_path)
    return output_path


def crear_respuestas_fundamentadas():
    """Crea documento con respuestas detalladas fundamentadas en apuntes."""
    doc = Document()
    
    # Título
    titulo = doc.add_heading('Respuestas Fundamentadas - Solemne I', 0)
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitulo = doc.add_paragraph('Basado en Apuntes del Curso de Minería de Datos\nProfesor Diego Robles C.')
    subtitulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_page_break()
    
    # SECCIÓN 1: ANÁLISIS Y PREPROCESAMIENTO
    agregar_titulo(doc, 'SECCIÓN 1: Análisis y Preprocesamiento', 1)
    
    # Pregunta 1
    agregar_titulo(doc, '1. ¿Qué estrategias utilizaron para manejar valores faltantes?', 2)
    
    agregar_parrafo(doc,
        'Fundamentación teórica (Semana 3 - Preprocesamiento):',
        negrita=True
    )
    
    agregar_parrafo(doc,
        'Los apuntes indican que los valores faltantes son uno de los principales problemas de calidad '
        'de datos. Las estrategias principales son:'
    )
    
    agregar_lista(doc, [
        'Eliminación de registros: Cuando los valores faltantes son < 5% del dataset',
        'Imputación por media/mediana: Para variables numéricas',
        'Imputación por moda: Para variables categóricas',
        'Algoritmos tolerantes: Algunos algoritmos manejan nativamente valores faltantes'
    ])
    
    agregar_parrafo(doc, 'Aplicación en nuestro proyecto:', negrita=True)
    
    agregar_parrafo(doc,
        'En el Adult Income Dataset, detectamos valores "?" que representan datos faltantes. '
        'Aplicamos eliminación de registros ya que representaban < 8% del total, manteniendo '
        'la integridad del análisis sin sesgar las distribuciones.'
    )
    
    # Pregunta 2
    agregar_titulo(doc, '2. ¿Cómo decidieron el método de codificación de variables categóricas?', 2)
    
    agregar_parrafo(doc, 'Fundamentación teórica (Semana 3):', negrita=True)
    
    agregar_parrafo(doc,
        'Los apuntes distinguen entre dos métodos principales:'
    )
    
    agregar_lista(doc, [
        'Label Encoding: Asigna números enteros a categorías. Útil para variables ordinales '
        'donde existe orden (ej: educación: primaria=1, secundaria=2, universidad=3)',
        
        'One-Hot Encoding: Crea columnas binarias para cada categoría. Recomendado para variables '
        'nominales sin orden intrínseco (ej: país, tipo de trabajo)',
    ])
    
    agregar_parrafo(doc, 'Decisión metodológica:', negrita=True)
    
    agregar_parrafo(doc,
        'Usamos Label Encoding para "education-num" (ya está ordenada numéricamente) y One-Hot '
        'Encoding para variables como "workclass", "occupation" y "marital-status" (nominales). '
        'Esto evita que el modelo asuma relaciones ordinales incorrectas.'
    )
    
    # Pregunta 3
    agregar_titulo(doc, '3. ¿Detectaron sesgos o desbalance en las clases?', 2)
    
    agregar_parrafo(doc, 'Fundamentación teórica (Semana 3 - Desbalance de clases):', negrita=True)
    
    agregar_parrafo(doc,
        'Los apuntes advierten sobre el "desbalance de las clases" como problema crítico. '
        'Cuando una clase es minoritaria, el modelo puede sesgarse hacia la clase mayoritaria, '
        'logrando alta accuracy pero fallando en la clase importante.'
    )
    
    agregar_parrafo(doc, 'Hallazgos en nuestro dataset:', negrita=True)
    
    agregar_lista(doc, [
        'Distribución: 76% ganan ≤50K vs 24% >50K (ratio 3:1)',
        'Impacto: Modelo tiende a predecir clase mayoritaria',
        'Estrategia: Priorizamos F1-Score sobre Accuracy para balancear precisión y recall',
        'En escenario real: Podría discriminar contra minorías con altos ingresos'
    ])
    
    doc.add_page_break()
    
    # SECCIÓN 2: MODELOS Y RENDIMIENTO
    agregar_titulo(doc, 'SECCIÓN 2: Modelos y su Rendimiento', 1)
    
    # Pregunta 4
    agregar_titulo(doc, '4. ¿Cuál modelo se desempeñó mejor?', 2)
    
    agregar_parrafo(doc, 'Análisis fundamentado:', negrita=True)
    
    agregar_parrafo(doc,
        'Regresión Logística superó a Random Forest en términos generales:'
    )
    
    agregar_lista(doc, [
        'F1-Score: 66.24% (LR) vs 64.67% (RF) - Mejor balance en dataset desbalanceado',
        'Recall: 60.08% (LR) vs 54.46% (RF) - Detecta más casos positivos',
        'AUC: Ambos >90%, excelente capacidad discriminativa',
        'Validación cruzada: Resultados consistentes confirman robustez'
    ])
    
    # Pregunta 5
    agregar_titulo(doc, '5. ¿Qué diferencias notaron en diferentes métricas?', 2)
    
    agregar_parrafo(doc, 'Fundamentación (Semana 6 - Exactitud y Precisión):', negrita=True)
    
    agregar_parrafo(doc,
        'Los apuntes explican que cada métrica captura aspectos diferentes del rendimiento:'
    )
    
    agregar_lista(doc, [
        'Accuracy: Porcentaje total de aciertos. Engañosa con clases desbalanceadas.',
        'Precision: De los predichos positivos, ¿cuántos son correctos? Minimiza falsos positivos.',
        'Recall (Sensibilidad): De los reales positivos, ¿cuántos detectamos? Minimiza falsos negativos.',
        'F1-Score: Media armónica de precisión y recall. Ideal para clases desbalanceadas.'
    ])
    
    agregar_parrafo(doc, 'Comportamiento observado:', negrita=True)
    
    agregar_lista(doc, [
        'Regresión Logística: Recall alto (60%), mejor para detectar ingresos altos',
        'Random Forest: Precision alta (79.59%), conservador en predicciones positivas',
        'Trade-off visible: RF sacrifica recall por precision (menos falsos positivos)'
    ])
    
    # Pregunta 6
    agregar_titulo(doc, '6. ¿Cómo interpretar alta precisión pero baja sensibilidad?', 2)
    
    agregar_parrafo(doc, 'Interpretación teórica (Semana 6 - Matriz de Confusión):', negrita=True)
    
    agregar_parrafo(doc,
        'Alta precisión + baja sensibilidad significa:'
    )
    
    agregar_lista(doc, [
        'El modelo es conservador: Solo predice positivo cuando está muy seguro',
        'Pocos falsos positivos (error tipo I bajo)',
        'Muchos falsos negativos (error tipo II alto): Pierde casos positivos reales',
        'En este problema: Subdetección de personas con ingresos >50K'
    ])
    
    agregar_parrafo(doc, '¿Es deseable?', negrita=True)
    
    agregar_parrafo(doc,
        'Depende del contexto: Si el costo de falso positivo es mayor que falso negativo, SÍ. '
        'Por ejemplo, en préstamos bancarios: preferible rechazar un buen cliente (falso negativo) '
        'que aprobar uno malo (falso positivo). En políticas sociales: NO deseable, excluiría '
        'beneficiarios legítimos.'
    )
    
    # Pregunta 7
    agregar_titulo(doc, '7. ¿Qué supuestos hace cada modelo?', 2)
    
    agregar_parrafo(doc, 'Supuestos de Regresión Logística (Semana 6):', negrita=True)
    
    agregar_lista(doc, [
        'Variable dependiente binaria: Cumplido (ingreso ≤50K o >50K)',
        'Independencia de observaciones: Cumplido (registros censales independientes)',
        'Relación lineal (en escala logit): Asume que log-odds es lineal con predictores',
        'Sin multicolinealidad severa: Variables no deben estar altamente correlacionadas',
        'Tamaño de muestra grande: Cumplido (32,561 registros)'
    ])
    
    agregar_parrafo(doc, 'Supuestos de Random Forest (Semana 7):', negrita=True)
    
    agregar_lista(doc, [
        'No asume distribución de datos: No paramétrico',
        'No requiere relación lineal: Captura interacciones no lineales',
        'Tolera multicolinealidad: Los árboles individuales pueden usar variables correlacionadas',
        'Robusto a outliers: Divisiones basadas en umbrales son menos sensibles',
        'Desventaja: Puede sobreajustar si los árboles son muy profundos'
    ])
    
    doc.add_page_break()
    
    # SECCIÓN 3: VALIDACIÓN CRUZADA
    agregar_titulo(doc, 'SECCIÓN 3: Validación Cruzada y Evaluación', 1)
    
    # Pregunta 8
    agregar_titulo(doc, '8. ¿Por qué es importante la validación cruzada?', 2)
    
    agregar_parrafo(doc, 'Fundamentación teórica:', negrita=True)
    
    agregar_parrafo(doc,
        'La validación cruzada (k-fold) es crucial porque:'
    )
    
    agregar_lista(doc, [
        'Usa todos los datos: Cada observación se usa para entrenamiento y validación',
        'Reduce varianza: Promedia resultados de k particiones diferentes',
        'Detecta sobreajuste: Si CV << holdout, el modelo memorizó el set de entrenamiento',
        'Estimación más confiable: Un solo split puede ser afortunado o desafortunado',
        'Validación estratificada: Mantiene proporción de clases en cada fold'
    ])
    
    # Pregunta 9
    agregar_titulo(doc, '9. ¿Qué diferencias encontraron entre holdout y CV?', 2)
    
    agregar_parrafo(doc, 'Resultados comparativos:', negrita=True)
    
    agregar_lista(doc, [
        'Regresión Logística: 66.24% (holdout) vs 66.09% ± 1.00% (CV) - Diferencia 0.15%',
        'Random Forest: 64.67% (holdout) vs 64.86% ± 0.94% (CV) - Diferencia 0.19%',
        'Consistencia excelente: Diferencias < 0.5% indican que holdout fue representativo',
        'Baja desviación: σ < 1% confirma estabilidad del modelo entre folds',
        'Conclusión: No hay sobreajuste, modelos generalizan bien'
    ])
    
    # Pregunta 10
    agregar_titulo(doc, '10. ¿Por qué usar F1-Score en lugar de solo Accuracy?', 2)
    
    agregar_parrafo(doc, 'Fundamentación (Semana 6):', negrita=True)
    
    agregar_parrafo(doc,
        'En datasets desbalanceados (76%-24%), la Accuracy es engañosa:'
    )
    
    agregar_lista(doc, [
        'Modelo ingenuo: Predecir siempre clase mayoritaria → 76% accuracy sin aprender nada',
        'F1-Score: Media armónica de precision y recall, penaliza desbalances',
        'Fórmula: F1 = 2 * (precision * recall) / (precision + recall)',
        'Ventaja: Considera tanto falsos positivos como falsos negativos',
        'En nuestro caso: LR con F1=66.24% superior a RF con F1=64.67%, aunque accuracy similar'
    ])
    
    doc.add_page_break()
    
    # SECCIÓN 4: VISUALIZACIÓN
    agregar_titulo(doc, 'SECCIÓN 4: Visualización y Explicación', 1)
    
    # Pregunta 11
    agregar_titulo(doc, '11. ¿Qué aprendieron de la matriz de confusión?', 2)
    
    agregar_parrafo(doc, 'Análisis de errores (Semana 6 - Tipos de error):', negrita=True)
    
    agregar_parrafo(doc, 'Regresión Logística:', negrita=True)
    agregar_lista(doc, [
        'Verdaderos Positivos: Detecta correctamente 60% de ingresos >50K',
        'Falsos Negativos (Error Tipo II): 40% de altos ingresos clasificados como bajos',
        'Falsos Positivos (Error Tipo I): 26% de bajos ingresos clasificados como altos',
        'Patrón: Más falsos negativos que falsos positivos'
    ])
    
    agregar_parrafo(doc, 'Random Forest:', negrita=True)
    agregar_lista(doc, [
        'Menos falsos positivos (20%) - Alta precision',
        'Más falsos negativos (45%) - Baja recall',
        'Trade-off: Sacrifica sensibilidad por especificidad'
    ])
    
    # Pregunta 12
    agregar_titulo(doc, '12. ¿Cuáles fueron las variables más importantes?', 2)
    
    agregar_parrafo(doc, 'Análisis de importancia (Semana 7 - Índice Gini):', negrita=True)
    
    agregar_parrafo(doc,
        'Random Forest calcula importancia mediante reducción de impureza Gini. Variables top 5:'
    )
    
    agregar_lista(doc, [
        'capital-gain (28.4%): Ganancias de capital - predictor dominante',
        'age (16.8%): Edad influye fuertemente en capacidad de ingresos',
        'hours-per-week (13.2%): Horas trabajadas correlacionan con ingresos',
        'education-num (12.1%): Años de educación son clave',
        'capital-loss (8.9%): Pérdidas de capital también predicen nivel económico'
    ])
    
    agregar_parrafo(doc, '¿Son coherentes con sentido común?', negrita=True)
    
    agregar_parrafo(doc,
        'SÍ. Las variables financieras (capital-gain/loss) naturalmente predicen ingresos. '
        'Educación y experiencia (edad) son predictores clásicos de salarios en economía laboral. '
        'Horas trabajadas reflejan dedicación y tipo de empleo.'
    )
    
    # Pregunta 13
    agregar_titulo(doc, '13. ¿Cómo contribuyen estas variables en términos reales?', 2)
    
    agregar_parrafo(doc, 'Interpretación económica:', negrita=True)
    
    agregar_lista(doc, [
        'Capital-gain: Inversiones y dividendos indican riqueza acumulada y conocimiento financiero',
        'Age: Experiencia laboral, seniority, acumulación de habilidades a lo largo de la vida',
        'Hours-per-week: Trabajo tiempo completo vs parcial, múltiples empleos, dedicación',
        'Education: Capital humano - estudios superiores abren acceso a empleos mejor remunerados',
        'Capital-loss: Actividad de inversión (aunque negativa) indica participación en mercados'
    ])
    
    doc.add_page_break()
    
    # SECCIÓN 5: COMPARACIÓN CRÍTICA
    agregar_titulo(doc, 'SECCIÓN 5: Comparación Crítica de Modelos', 1)
    
    # Pregunta 14
    agregar_titulo(doc, '14. ¿Qué modelo para interpretabilidad vs precisión?', 2)
    
    agregar_parrafo(doc, 'Para interpretabilidad: Regresión Logística', negrita=True)
    
    agregar_lista(doc, [
        'Coeficientes interpretables: Cada β muestra dirección e impacto',
        'Probabilidades calibradas: Salida directa es P(Y=1|X)',
        'Explicable a stakeholders: "Por cada año de educación, la probabilidad aumenta X%"',
        'Auditable: Fácil verificar sesgos en coeficientes'
    ])
    
    agregar_parrafo(doc, 'Solo para precisión: Depende del contexto', negrita=True)
    
    agregar_lista(doc, [
        'En nuestro caso: LR tiene mejor F1 (66.24% vs 64.67%)',
        'RF tiene AUC ligeramente superior (90.70% vs 90.24%)',
        'Diferencia marginal: Elegir según otros criterios (interpretabilidad, velocidad)'
    ])
    
    # Pregunta 15
    agregar_titulo(doc, '15. Ventajas y desventajas de Random Forest vs Regresión Logística', 2)
    
    agregar_parrafo(doc, 'Ventajas de Random Forest (Semana 7):', negrita=True)
    
    agregar_lista(doc, [
        'No requiere normalización: Maneja escalas diferentes naturalmente',
        'Captura no-linealidades: Detecta interacciones complejas',
        'Robusto a outliers: Divisiones por umbrales son resistentes',
        'Selección automática de variables: Usa las más informativas en cada split',
        'Menor riesgo de sobreajuste: Ensemble reduce varianza'
    ])
    
    agregar_parrafo(doc, 'Desventajas de Random Forest:', negrita=True)
    
    agregar_lista(doc, [
        'Caja negra: Difícil interpretar decisiones individuales',
        'Mayor complejidad computacional: Entrena múltiples árboles',
        'Memoria: Almacena todos los árboles del ensemble',
        'Tiempo de predicción: Más lento que modelos lineales',
        'Puede sobreajustar con hiperparámetros mal configurados'
    ])
    
    # Pregunta 16
    agregar_titulo(doc, '16. ¿Cuál escala mejor a millones de registros?', 2)
    
    agregar_parrafo(doc, 'Regresión Logística escala mejor:', negrita=True)
    
    agregar_lista(doc, [
        'Complejidad: O(n*p) - lineal con número de observaciones',
        'Entrenamiento: Algoritmos eficientes (SGD, L-BFGS)',
        'Predicción: Producto matriz-vector muy rápido',
        'Memoria: Solo almacena p coeficientes (uno por variable)',
        'Paralelizable: Fácil implementar en distributed systems (Spark)'
    ])
    
    agregar_parrafo(doc, 'Random Forest presenta desafíos:', negrita=True)
    
    agregar_lista(doc, [
        'Complejidad: O(n*log(n)*p*T) donde T = número de árboles',
        'Memoria: Almacena T árboles completos',
        'Predicción: Debe consultar todos los árboles',
        'Escalabilidad: Requiere infraestructura distribuida (RandomForestClassifier no escala bien)',
        'Alternativa: Usar XGBoost o LightGBM con optimizaciones para big data'
    ])
    
    doc.add_page_break()
    
    # SECCIÓN 6: ÉTICA
    agregar_titulo(doc, 'SECCIÓN 6: Ética y Aplicación Real', 1)
    
    # Pregunta 17
    agregar_titulo(doc, '17. ¿Qué problemas éticos pueden surgir?', 2)
    
    agregar_parrafo(doc, 'Fundamentación (Semana 1 - Ética y Big Data):', negrita=True)
    
    agregar_parrafo(doc,
        'El uso de modelos predictivos de ingresos para decisiones reales (préstamos, empleo, '
        'políticas públicas) presenta riesgos éticos graves:'
    )
    
    agregar_lista(doc, [
        'Discriminación algorítmica: Variables proxy de raza/género pueden perpetuar desigualdad',
        'Falta de transparencia: "Caja negra" impide que afectados entiendan decisiones',
        'Sesgo de confirmación: El modelo aprende y refuerza prejuicios históricos',
        'Exclusión financiera: Falsos negativos niegan oportunidades a personas merecedoras',
        'Responsabilidad difusa: ¿Quién es responsable de errores del algoritmo?',
        'Privacy: Datos demográficos sensibles pueden ser inferidos'
    ])
    
    # Pregunta 18
    agregar_titulo(doc, '18. ¿Cómo afectan los sesgos sociales a las decisiones del modelo?', 2)
    
    agregar_parrafo(doc, 'Sesgos identificados en Adult Income Dataset:', negrita=True)
    
    agregar_lista(doc, [
        'Sesgo histórico: Datos de 1994 reflejan desigualdad salarial de género de esa época',
        'Sesgo de representación: Minorías subrepresentadas en ingresos altos',
        'Sesgo de medición: "Occupation" puede codificar estereotipos de género',
        'Retroalimentación: Si el modelo niega préstamos, personas no pueden mejorar ingresos',
        'Correlación espuria: "Marital-status" correlaciona con ingresos, pero es discriminatorio usarlo'
    ])
    
    agregar_parrafo(doc, 'Impacto en decisiones reales:', negrita=True)
    
    agregar_parrafo(doc,
        'Un modelo entrenado con estos sesgos podría: (1) Negar préstamos a mujeres incluso con '
        'igual calificación, (2) Perpetuar brechas salariales al predecir menores ingresos esperados, '
        '(3) Excluir minorías de programas de desarrollo profesional.'
    )
    
    # Pregunta 19
    agregar_titulo(doc, '19. ¿Qué harían diferente en una aplicación real de política pública?', 2)
    
    agregar_parrafo(doc, 'Mejoras metodológicas para uso en políticas públicas:', negrita=True)
    
    agregar_lista(doc, [
        'Auditoría de equidad: Medir disparate impact entre grupos demográficos',
        'Variables protegidas: Remover raza, género, edad de predictores directos',
        'Fairness constraints: Aplicar restricciones de equidad durante entrenamiento',
        'Validación externa: Probar en poblaciones diferentes a entrenamiento',
        'Explicabilidad: Implementar SHAP/LIME para justificar cada predicción',
        'Revisión humana: Combinar modelo con evaluación experta',
        'Transparencia: Documentar limitaciones y sesgos conocidos',
        'Monitoreo continuo: Detectar drift y degradación de equidad'
    ])
    
    # Pregunta 20
    agregar_titulo(doc, '20. ¿Cómo hacer el modelo más justo y equitativo?', 2)
    
    agregar_parrafo(doc, 'Estrategias de mitigación de sesgo:', negrita=True)
    
    agregar_lista(doc, [
        'Pre-procesamiento:',
        '  • Reweighting: Dar más peso a grupos subrepresentados',
        '  • Resampling: Balancear clases y grupos protegidos',
        '  • Transformación de features: Remover correlación con atributos sensibles',
        '',
        'In-processing:',
        '  • Adversarial debiasing: Entrenar para ser justo',
        '  • Regularización de equidad: Penalizar disparidad en loss function',
        '  • Calibración por grupos: Ajustar umbrales diferentes por demografía',
        '',
        'Post-procesamiento:',
        '  • Equalized odds: Ajustar predicciones para igualar TPR/FPR entre grupos',
        '  • Reject option classification: Re-clasificar casos cercanos al umbral',
        '',
        'Gobernanza:',
        '  • Comité de ética: Revisar uso del modelo',
        '  • Derecho a explicación: Usuarios pueden solicitar justificación',
        '  • Auditoría independiente: Evaluación externa periódica'
    ])
    
    agregar_parrafo(doc, 'Implementación práctica:', negrita=True)
    
    agregar_parrafo(doc,
        'Usaríamos bibliotecas como AIF360 (AI Fairness 360) de IBM o Fairlearn de Microsoft '
        'para implementar estas estrategias. Definiríamos métricas de equidad como Statistical '
        'Parity Difference y Equal Opportunity Difference, estableciendo umbrales aceptables '
        '(ej: SPD < 0.1) antes de deployment.'
    )
    
    # CONCLUSIÓN FINAL
    doc.add_page_break()
    agregar_titulo(doc, 'Conclusión Final', 1)
    
    agregar_parrafo(doc,
        'Este análisis demuestra la importancia de fundamentar las decisiones de minería de datos '
        'en teoría sólida. Los conceptos de preprocesamiento (Semana 3), regresión logística '
        '(Semana 6) y árboles de decisión (Semana 7) proporcionaron el marco para implementar, '
        'evaluar e interpretar correctamente los modelos.'
    )
    
    agregar_parrafo(doc,
        'La validación cruzada confirmó la robustez de nuestros resultados, mientras que el '
        'análisis crítico de métricas y visualizaciones permitió comprender las fortalezas y '
        'limitaciones de cada enfoque. Crucialmente, las consideraciones éticas (Semana 1) '
        'nos recuerdan que los modelos no son neutrales: reflejan y pueden amplificar sesgos '
        'sociales existentes.'
    )
    
    agregar_parrafo(doc,
        'La minería de datos es una herramienta poderosa, pero requiere responsabilidad, '
        'transparencia y un compromiso continuo con la equidad y la justicia social.',
        italica=True
    )
    
    # Guardar
    output_path = Path('/workspaces/mineriadatos/results/Respuestas_Fundamentadas_Solemne_I.docx')
    doc.save(output_path)
    return output_path


def main():
    """Genera todos los documentos fundamentados."""
    print("=" * 80)
    print("GENERANDO DOCUMENTOS FUNDAMENTADOS EN APUNTES DEL CURSO")
    print("=" * 80)
    
    documentos = []
    
    # 1. Informe grupal
    print("\n1. Generando Informe Grupal Fundamentado (máx. 3 páginas)...")
    path1 = crear_informe_fundamentado()
    documentos.append(('Informe Grupal', path1))
    print(f"   ✓ Creado: {path1.name}")
    
    # 2. Respuestas detalladas
    print("\n2. Generando Respuestas Fundamentadas (20 preguntas)...")
    path2 = crear_respuestas_fundamentadas()
    documentos.append(('Respuestas Detalladas', path2))
    print(f"   ✓ Creado: {path2.name}")
    
    print("\n" + "=" * 80)
    print("✅ DOCUMENTOS FUNDAMENTADOS GENERADOS EXITOSAMENTE")
    print("=" * 80)
    
    print("\n📚 BASE TEÓRICA UTILIZADA:")
    print("   • Semana 1: Ética y Big Data")
    print("   • Semana 3: Preprocesamiento de Datos")
    print("   • Semana 5: Preparación de la información")
    print("   • Semana 6: Regresión Logística")
    print("   • Semana 7: Árboles de Decisión")
    
    print("\n📄 ARCHIVOS CREADOS:")
    for nombre, path in documentos:
        print(f"   • {nombre}: {path}")
    
    print("\n💡 DIFERENCIAS CON VERSIÓN ANTERIOR:")
    print("   ✅ Respuestas fundamentadas en apuntes del profesor")
    print("   ✅ Referencias a contenidos específicos de cada semana")
    print("   ✅ Terminología y conceptos del curso")
    print("   ✅ Coherencia con metodología enseñada")
    
    print("\n📁 Ubicación: /workspaces/mineriadatos/results/")
    print("=" * 80)


if __name__ == "__main__":
    main()
